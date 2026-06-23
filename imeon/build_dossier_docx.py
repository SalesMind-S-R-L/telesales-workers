#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY = RGBColor(0x1F,0x38,0x64)
doc = Document()
st = doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(10.5)

def h1(t):
    p=doc.add_heading(t,level=1)
    for r in p.runs: r.font.color.rgb=NAVY
def h2(t):
    p=doc.add_heading(t,level=2)
    for r in p.runs: r.font.color.rgb=NAVY
def para(t,bold=False,italic=False,size=None):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=bold; r.italic=italic
    if size: r.font.size=Pt(size)
    return p
def bullet(t):
    doc.add_paragraph(t,style='List Bullet')

# Cover
title=doc.add_paragraph(); title.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=title.add_run('DOSSIER MERCATO IMMOBILIARE'); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=NAVY
sub=doc.add_paragraph(); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=sub.add_run('Firenze Sud-Ovest + Scandicci'); r.font.size=Pt(14); r.font.color.rgb=NAVY
s2=doc.add_paragraph(); s2.alignment=WD_ALIGN_PARAGRAPH.CENTER
s2.add_run("Per acquisizione immobili - \"L'agenzia alle torri\" (ex Tecnocasa)").italic=True
s3=doc.add_paragraph(); s3.alignment=WD_ALIGN_PARAGRAPH.CENTER
s3.add_run('Aggiornamento: giugno 2026  |  Zone: Isolotto, Legnaia, Soffiano, Ponte a Greve, Firenze Q4, Scandicci').font.size=Pt(9)
para("Nota metodologica: i valori al mq incrociano asking price (Immobiliare.it via aggregatori), quotazioni notarili OMI (Agenzia Entrate) e report Ufficio Studi Tecnocasa. Gli asking price sono ~7% sopra i prezzi di rogito effettivi. I valori OMI sono i piu bassi perche basati su atti registrati.",italic=True,size=9)

def table(headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.style='Light Grid Accent 1'
    hc=t.rows[0].cells
    for i,hh in enumerate(headers):
        hc[i].text=''; run=hc[i].paragraphs[0].add_run(hh); run.bold=True; run.font.size=Pt(9)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=''; rr=cells[i].paragraphs[0].add_run(str(v)); rr.font.size=Pt(9)
    return t

h1('1. Prezzi al mq - Vendita')
table(['Zona','Medio EUR/mq','Range EUR/mq','Trend 12 mesi','Fonte'],[
 ['Firenze citta','4.719-4.737','3.432-6.364','+5,4% (apr 26 vs 25)','Immobiliare.it / catonecasa'],
 ['Firenze Sud (Q4)','4.696','4.450-4.900','+18,65%','catonecasa / Immobiliare.it'],
 ['Legnaia-Soffiano','3.946','-','+16,57%','catonecasa / Immobiliare.it'],
 ['Isolotto','3.732','-','+15,53%','Immobiliare.it / catonecasa'],
 ['Scandicci (asking)','3.382-3.538','2.809-3.736','+4,3% / +7%','Immobiliare.it'],
 ['Scandicci (OMI rogiti)','2.670','appart. 2.510 / case 2.723','+6,04%','OMI Agenzia Entrate'],
])
bullet('Il quadrante sud-ovest cresce piu della media citta (+15-18% vs +5,4%): zona "in recupero". Isolotto/Legnaia restano sotto la media citta (3.700-3.950 vs 4.737) -> accessibili, forte domanda.')
bullet('Sub-zone Scandicci: max Mosciano-Casignano-Giogoli 3.736 EUR/mq; min San Vincenzo a Torri / S. Maria a Marciola 2.809 EUR/mq. Picco storico giugno 2025.')
bullet('Previsione Firenze 2026: oltre 5.000 EUR/mq.')

h1('2. Affitti (EUR/mq/mese)')
table(['Zona','EUR/mq/mese','Note'],[
 ['Firenze citta','20,75','previsione 2026 verso 23'],
 ['Isolotto (asking)','15,53','zona piu economica di Firenze per l\'affitto'],
 ['Isolotto (OMI C15)','7,1-9,2','valori ufficiali Agenzia Entrate'],
])
para('Riferimenti reali: monolocale 29 mq Via Pisana (Legnaia) 650 EUR/mese; trilocale zona sud ~800 EUR/mese. Isolotto = ingresso piu economico della citta -> alta rotazione -> bacino di proprietari che prima o poi vendono.')

h1('3. Domanda vs Offerta - mercato "tirato"')
bullet('Tempi di vendita Firenze citta: ~97 giorni (gen 2025, Tecnocasa); provincia tra le piu rapide d\'Italia. II sem 2025: assorbimento ~6 mesi. Hinterland 117 gg.')
bullet('Sconto medio sul richiesto a Firenze: 7% (vs 11,2% media nazionale) -> a Firenze si tratta poco = mercato del venditore.')
bullet('Offerta in calo + domanda solida -> la scarsita sostiene i valori e riduce i tempi. Compravendite Italia I sem 2025 +9,5%.')
para('Implicazione acquisizione: il prodotto da vendere e scarso -> chi ACQUISISCE l\'incarico per primo e il collo di bottiglia. Vende chiunque; il valore e mettere le mani sull\'incarico, meglio se in esclusiva.',bold=True)
para('Esclusiva: agenzie top 85% portafoglio in esclusiva; 1 esclusiva = 5-8 incarichi non esclusivi per ROI.')

h1('4. Ciclo di ricambio casa (claim "7-8 anni")')
para('Il numero secco "7-8 anni" NON e confermato dalle fonti pubbliche - usare con cautela. Dati certi:')
bullet('~30% degli italiani cambia casa almeno una volta ogni 5 anni (Weplaza).')
bullet('Nuclei che cambiano casa: da 797.000 (2001) a 1.617.000 (2022) -> dal 3,7% al 6,4%/anno.')
bullet('4.430 nuclei/giorno cambiano indirizzo. 82% vive in casa di proprieta (ISTAT dic 2025).')
para('Riformulazione difendibile: "una quota crescente di famiglie cambia casa ogni 5-8 anni" - non citare il 7-8 come dato ufficiale.',bold=True)

h1('5. Profilo acquirente tipico (Tecnocasa Firenze 2025)')
bullet('81,5% acquista abitazione principale; 15,2% investimento; 3,3% casa vacanza.')
bullet('74,3% coppie e famiglie; single in crescita 22,8% -> 25,7%.')
bullet('Investimento: Firenze la % piu alta d\'Italia 22,5%; 77,1% degli investimenti in contanti.')
bullet('Stranieri in calo: 5,7%. Tipologia piu richiesta: trilocale (2 camere) 32,7%, alla pari col quadrilocale.')
para('Chi cerca casa e probabilmente ha casa da vendere: la coppia/famiglia in upgrade da bilo/trilocale a trilo/quadrilocale (per figli) - ha quasi sempre un immobile da liberare. Nel sud-ovest (zona accessibile, prima casa) il trade-up e frequente.',bold=True)

h1('6. Concorrenza (Isolotto / Legnaia / Scandicci)')
bullet('Densita altissima: ~190 agenzie in Isolotto-Legnaia (Q4), 32 a Scandicci -> l\'acquisizione e guerra di velocita e relazione.')
bullet('Franchising: Tecnocasa (il cliente e ex-Tecnocasa, conosce il metodo), Gabetti, Toscano, Grimaldi/L\'Immobiliare.com, RE/MAX (Isolotto-Talenti), Italiana Immobiliare (Soffiano, Scandicci).')
bullet('Differenziale: non la vendita (il prodotto si vende da solo in 3-6 mesi) ma chi prende l\'incarico per primo e in esclusiva.')

h1('7. Leve di acquisizione 2026')
bullet('Valutazione professionale (analisi di mercato sui prezzi reali), non "gratis e generica" (attira perditempo).')
bullet('Scouting annunci privati (FSBO): proprietari motivati con prezzo/tempi sub-ottimali, target naturale per esclusiva.')
bullet('"Deve prima vendere": chi compra (trade-up famiglie 74%) ha quasi sempre un immobile da vendere -> ogni lead-acquirente e anche lead-venditore.')
bullet('Direttiva Case Green: classi energetiche basse perdono valore -> urgenza per chi e indeciso.')
bullet('Esclusiva come obiettivo (ROI 5-8x).')

h1('8. Takeaway operativi (insight -> uso nello script)')
tk=[
 ('Mercato del venditore (sconto 7%, 97 gg)','"In zona le case si vendono in meno di 3 mesi con sconti minimi. Se ha un immobile e il momento: vuole sapere quanto vale oggi?"'),
 ('Prezzi sud-ovest +15-18% in un anno','"Forse la sua casa vale molto piu di quanto pensa. Le faccio una valutazione aggiornata?"'),
 ('Offerta scarsa','"Pochissime case sul mercato, tanti acquirenti in attesa. Se mette in vendita ora ha la fila."'),
 ('Chi cerca casa deve prima vendere la sua','"Quando cercava casa, doveva anche venderne una? Quella che ha oggi e ancora sua?"'),
 ('Trilocale taglio piu richiesto (32,7%)','"I trilocali in zona vanno a ruba, ho gia acquirenti che li cercano."'),
 ('Isolotto zona piu economica -> alta rotazione','Targetizzare chi e in zona da 5-8 anni: tipico momento di upgrade.'),
 ('Valutazione professionale, non generica','"Le porto un\'analisi sui prezzi reali degli ultimi atti nella sua via, non una stima a caso."'),
 ('Direttiva Case Green come urgenza','"Con le nuove regole UE le classi basse rischiano di perdere valore."'),
 ('190 competitor -> velocita ed esclusiva','Chiudere su appuntamento fisico ravvicinato; puntare all\'esclusiva.'),
 ('Investimento cash alto (77%)','Per chi ha immobile a reddito: "Ho acquirenti liquidi pronti, lo valutiamo?"'),
]
table(['Insight chiave','Come usarlo in chiamata'],tk)

h1('Caveat sui dati')
bullet('Valori Immobiliare.it = asking price (offerta), reali ~7% sotto. OMI = da rogiti, piu conservativi.')
bullet('"Ciclo casa 7-8 anni" non verificabile come cifra secca.')
bullet('Dati Scandicci ad ago 2025, Firenze ad apr 2026 (sfasamento fisiologico tra fonti).')

doc.add_paragraph()
para('Fonti principali: Immobiliare.it, Idealista.it, catonecasa.it, OMI Agenzia delle Entrate (immobiliagest.it), Ufficio Studi Tecnocasa, Nomisma, 055firenze.it, La Nazione, ISTAT/ANSA, peracquisire.it, RealAdvisor, Gromia.',italic=True,size=8)

out='/Users/simocors/Desktop/telesales/imeon/Dossier_Mercato_IMEON.docx'
doc.save(out)
print('SALVATO',out)
