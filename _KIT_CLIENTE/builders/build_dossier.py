#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
BUILDER FISSO - Dossier di mercato (Word) cliente.
USO: python3 build_dossier.py /percorso/cartella_cliente
Legge: config.json (cliente/data) + dossier.json
Output: Dossier_Mercato_<CLIENTE>.docx

dossier.json:
{
 "titolo": "DOSSIER DI MERCATO",
 "sottotitolo": "...",
 "nota_metodo": "...",                         (opzionale, in corsivo sotto la cover)
 "sezioni": [
   {"titolo":"1. ...","paragrafi":["..."],"bullet":["..."],
    "tabella":{"headers":["..."],"rows":[["..",".."]]}}
 ],
 "takeaway": [["Insight chiave","Come usarlo in chiamata"], ...],
 "fonti": [["Area","https://..."], ...]
}
"""
import sys, os, json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

WD = sys.argv[1] if len(sys.argv) > 1 else '.'
def load(name, d=None):
    p = os.path.join(WD, name)
    return json.load(open(p, encoding='utf-8')) if os.path.exists(p) else d
config = load('config.json', {})
D = load('dossier.json', {})
import re
CLIENTE = config.get('cliente', 'CLIENTE'); SAFE = re.sub(r'[^A-Za-z0-9]+', '_', CLIENTE).strip('_')

NAVY = RGBColor(0x1F, 0x38, 0x64)
doc = Document(); st = doc.styles['Normal']; st.font.name = 'Calibri'; st.font.size = Pt(10.5)
def h1(t):
    p = doc.add_heading(t, level=1)
    for r in p.runs: r.font.color.rgb = NAVY
def para(t, bold=False, italic=False, size=None):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = bold; r.italic = italic
    if size: r.font.size = Pt(size)
    return p
def bullet(t): doc.add_paragraph(t, style='List Bullet')
def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Light Grid Accent 1'
    for i, hh in enumerate(headers):
        t.rows[0].cells[i].text = ''; run = t.rows[0].cells[i].paragraphs[0].add_run(str(hh)); run.bold = True; run.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ''; rr = cells[i].paragraphs[0].add_run(str(v)); rr.font.size = Pt(9)

# cover
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run(D.get('titolo', 'DOSSIER DI MERCATO')); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = NAVY
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = s.add_run(D.get('sottotitolo', '')); rr.font.size = Pt(13); rr.font.color.rgb = NAVY
s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
s2.add_run(f'Per {CLIENTE}  |  {config.get("data","")}').italic = True
if D.get('nota_metodo'): para(D['nota_metodo'], italic=True, size=9)

for sez in D.get('sezioni', []):
    h1(sez.get('titolo', ''))
    for p in sez.get('paragrafi', []): para(p)
    tb = sez.get('tabella')
    if tb and tb.get('headers'): table(tb['headers'], tb.get('rows', []))
    for b in sez.get('bullet', []): bullet(b)

if D.get('takeaway'):
    h1('Takeaway operativi (insight -> uso in chiamata)')
    table(['Insight chiave', 'Come usarlo in chiamata'], D['takeaway'])

if D.get('fonti'):
    h1('Fonti')
    for f in D['fonti']:
        para(' - '.join(str(x) for x in f), size=8)

out = os.path.join(WD, f'Dossier_Mercato_{SAFE}.docx')
doc.save(out); print('SALVATO', out)
