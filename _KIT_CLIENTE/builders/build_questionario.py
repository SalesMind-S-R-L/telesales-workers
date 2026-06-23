#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
BUILDER FISSO - Questionario onboarding (Word) cliente.
USO: python3 build_questionario.py /percorso/cartella_cliente
Legge: config.json + questions.json
Output: Questionario_<CLIENTE>.docx

REGOLA: massimo 15 domande totali, in ordine di priorita (il builder taglia a 15
se ne arrivano di piu, e numera in ordine). Sotto ogni domanda lo spazio "Risposta:".

questions.json:
{
 "titolo": "ONBOARDING",                         (opzionale)
 "come_usare": "...",
 "insight_settore": [ {"titolo":"1. ...","testo":"..."} ],      (Parte 1, ~5-6, non contano nel limite di 15)
 "macro_aree": [ {"nome":"MACRO A - ...","intro":"...","domande":[ {"titolo":"...","testo":"...","prio":1} ]} ],
 "kpi": [ {"titolo":"...","testo":"...","prio":1} ]              (le KPI sono domande, contano nel limite di 15)
}
"""
import sys, os, json, re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

WD = sys.argv[1] if len(sys.argv) > 1 else '.'
def load(name, d=None):
    p = os.path.join(WD, name)
    return json.load(open(p, encoding='utf-8')) if os.path.exists(p) else d
config = load('config.json', {}); Q = load('questions.json', {})
CLIENTE = config.get('cliente', 'CLIENTE'); SAFE = re.sub(r'[^A-Za-z0-9]+', '_', CLIENTE).strip('_')
REFERENTE = config.get('referente', '')
MAXQ = 15

NAVY = RGBColor(0x1F, 0x38, 0x64); GREY = RGBColor(0x55, 0x55, 0x55)
doc = Document(); st = doc.styles['Normal']; st.font.name = 'Calibri'; st.font.size = Pt(10.5)
def h1(t):
    p = doc.add_heading(t, level=1)
    for r in p.runs: r.font.color.rgb = NAVY
def h2(t):
    p = doc.add_heading(t, level=2)
    for r in p.runs: r.font.color.rgb = NAVY
def para(t, bold=False, italic=False, size=None):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = bold; r.italic = italic
    if size: r.font.size = Pt(size)
def bullet(t): doc.add_paragraph(t, style='List Bullet')
def rule(): para('_' * 60)

# raccogli tutte le domande con priorita, taglia a 15 in ordine
qlist = []
for area in Q.get('macro_aree', []):
    for d in area.get('domande', []):
        qlist.append({'area': area.get('nome', ''), 'titolo': d.get('titolo', ''), 'testo': d.get('testo', ''), 'prio': d.get('prio', 99)})
for d in Q.get('kpi', []):
    qlist.append({'area': 'KPI e modalita di lavoro', 'titolo': d.get('titolo', ''), 'testo': d.get('testo', ''), 'prio': d.get('prio', 99)})
qlist.sort(key=lambda x: x['prio'])
tagliate = max(0, len(qlist) - MAXQ)
qlist = qlist[:MAXQ]
# ri-raggruppa per area mantenendo l'ordine di comparsa
areas_order = []
for q in qlist:
    if q['area'] not in areas_order: areas_order.append(q['area'])

def render_q(code, q):
    p = doc.add_paragraph(); r = p.add_run(f'{code}. {q["titolo"]}'); r.bold = True; r.font.color.rgb = NAVY
    para(q['testo'])
    rp = doc.add_paragraph(); rr = rp.add_run('Risposta: '); rr.italic = True; rr.font.color.rgb = GREY
    doc.add_paragraph()

# cover
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run(f'{CLIENTE} - {Q.get("titolo","Onboarding")}'); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = NAVY
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = s.add_run('Insight di mercato + domande strutturate per chiudere i requisiti'); rr.font.size = Pt(13); rr.font.color.rgb = NAVY
s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
s2.add_run(f'Telesales  |  {config.get("data","")}' + (f'  |  Per {REFERENTE}' if REFERENTE else '')).italic = True
rule()

h2('Come usare questo documento')
para(Q.get('come_usare', 'Questo documento serve da agenda: (1) i punti chiave dalla nostra ricerca di mercato; (2) le domande operative per chiudere i requisiti; (3) KPI e cadenza. Puoi rispondere a voce in call o via email: sotto ogni domanda trovi lo spazio "Risposta:".'))

# Parte 1 - insight (non contano nel limite)
if Q.get('insight_settore'):
    rule(); h1('PARTE 1 - Insight settore')
    para('Sintesi dei punti chiave dalla nostra ricerca (dossier completo a parte). Ci allineano sulla stessa lettura del mercato.')
    for ins in Q['insight_settore']:
        para(ins.get('titolo', ''), bold=True); para(ins.get('testo', ''))

# Parti per macro area (max 15 domande totali, gia prioritizzate)
n = 0
letters = {}
for area in areas_order:
    rule(); h1(area)
    qs = [q for q in qlist if q['area'] == area]
    # prefisso lettera per area (A1, A2, ... / K1 per KPI)
    pref = 'K' if area.lower().startswith('kpi') else chr(ord('A') + areas_order.index(area))
    for i, q in enumerate(qs, 1):
        render_q(f'{pref}{i}', q)

rule()
chiusura = Q.get('chiusura', 'Le risposte ci permettono di chiudere i requisiti, mandarti la lista definitiva da approvare e partire con le chiamate. Grazie.')
para(chiusura, italic=True, size=9.5)
if tagliate:
    para(f'(Nota interna: {tagliate} domande a priorita piu bassa sono state rimandate al round successivo per restare entro le {MAXQ}.)', italic=True, size=8)

out = os.path.join(WD, f'Questionario_{SAFE}.docx')
doc.save(out); print('SALVATO', out, f'| domande incluse: {len(qlist)} (max {MAXQ}), rimandate: {tagliate}')
