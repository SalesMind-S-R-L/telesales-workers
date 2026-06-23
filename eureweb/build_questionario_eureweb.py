#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Questionario onboarding Eureweb / AD Lab - modello Giglioli (domande da far rispondere a Giulia)."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY = RGBColor(0x1F,0x38,0x64)
GREY = RGBColor(0x55,0x55,0x55)
doc = Document()
st = doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(10.5)

def h1(t):
    p=doc.add_heading(t,level=1)
    for r in p.runs: r.font.color.rgb=NAVY
def h2(t):
    p=doc.add_heading(t,level=2)
    for r in p.runs: r.font.color.rgb=NAVY
def para(t,bold=False,italic=False,size=None,color=None):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=bold; r.italic=italic
    if size: r.font.size=Pt(size)
    if color: r.font.color.rgb=color
    return p
def bullet(t):
    doc.add_paragraph(t,style='List Bullet')
def rule():
    p=doc.add_paragraph(); r=p.add_run('_'*60); r.font.color.rgb=GREY

def q(code, titolo, testo):
    p=doc.add_paragraph(); r=p.add_run(f'{code}. {titolo}'); r.bold=True; r.font.color.rgb=NAVY
    para(testo)
    rp=doc.add_paragraph(); rr=rp.add_run('Risposta: '); rr.italic=True; rr.font.color.rgb=GREY
    doc.add_paragraph()  # spazio per scrivere

# ===== COVER =====
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run('EUREWEB / AD LAB - Onboarding'); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=NAVY
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=s.add_run('Insight di mercato + domande strutturate per chiudere i requisiti del target'); r.font.size=Pt(13); r.font.color.rgb=NAVY
s2=doc.add_paragraph(); s2.alignment=WD_ALIGN_PARAGRAPH.CENTER
s2.add_run('Telesales  |  15 giugno 2026  |  Per Giulia Rizzi').italic=True
rule()

# ===== COME USARE =====
h2('Come usare questo documento')
para('Giulia, questo documento serve da agenda per allineare e chiudere la fase di setup. Tre parti: (1) i punti chiave emersi dalla nostra ricerca di mercato - 5 minuti di lettura, ci danno la stessa lettura del mercato; (2) le domande operative divise per area, servono a chiudere i requisiti del target e a validare la lista; (3) KPI e cadenza di lavoro, per fissare gli accordi pratici. Puoi rispondere a voce nella prossima call oppure anticiparci qualcosa via email: sotto ogni domanda trovi lo spazio "Risposta:".')
para('Riferimenti gia pronti: Master_Eureweb_ADLab.xlsx (427 prospect verificati uno a uno per investimento in adv/contenuti, con dimensione e accessibilita del decisore) e Dossier_Mercato_Eureweb.docx (ricerca di mercato completa con fonti).')

# ===== INDICE =====
h2('Indice')
bullet('Parte 1 - Insight settore: cosa abbiamo trovato (6 punti chiave)')
bullet('Parte 2A - MACRO A: Requisiti del target e ICP (6 domande)')
bullet('Parte 2B - MACRO B: Lista, scouting e messaggio (5 operative + 5 non-ovvie)')
bullet('Parte 3 - KPI e modalita di lavoro (4 domande)')
rule()

# ===== PARTE 1 =====
h1('PARTE 1 - Insight settore')
para('Sintesi dei punti chiave dalla nostra ricerca (dossier completo: Dossier_Mercato_Eureweb.docx). Una volta condivisi, ci allineano sulla stessa lettura del mercato e rendono coerenti le scelte operative.')

para('1. Il video e il cuore dell\'advertising italiano', bold=True)
para('Il video pesa circa il 43% del totale advertising (mercato dell\'ordine di 4,7 mld EUR; fonti PoliMi / Confindustria Radio TV). Take operativo: AD Lab produce esattamente l\'asset piu richiesto dal mercato - lo diciamo in apertura, non come "novita" ma come "il formato su cui i brand stanno gia spostando i budget".')

para('2. L\'AI nei contenuti non e piu sperimentazione', bold=True)
para('Mercato AI in Italia intorno a 1,8 mld EUR e in forte crescita; adozione della GenAI in netto aumento nelle grandi imprese; casi enterprise concreti (Coca-Cola ha generato decine di migliaia di clip con GenAI nel 2025). Take operativo: il pitch non e "rischiamo qualcosa di nuovo" ma "i grandi brand lo stanno gia facendo, vi mostriamo come farlo meglio".')

para('3. Costo: tradizionale contro AI', bold=True)
para('Uno shooting o uno spot tradizionale ha costi elevati; la produzione AI consente la stessa resa a una frazione del costo e in tempi piu rapidi. Take operativo: e una leva fortissima, MA in chiamata NON si fanno prezzi: la usiamo come concetto ("stessa qualita, tempi e costi piu efficienti"), il numero lo gestisci tu in appuntamento.')

para('4. Chi decide i contenuti nei grandi brand', bold=True)
para('Il Direttore Marketing guida la scelta dei partner; sotto di lui Digital Marketing Manager e PR/Comunicazione; sopra CEO/CdA che sbloccano i budget. I cicli enterprise sono lunghi e multi-touch (settimane/mesi, piu decisori). Take operativo: lo scouting LinkedIn punta a marketing + C-level; nelle multinazionali si lavora sull\'entita italiana.')

para('5. I 6 settori della lista sono i veri top spender', bold=True)
para('Secondo i dati Nielsen, i settori che investono di piu in pubblicita in Italia includono automotive, food & beverage, retail/GDO, finance/assicurazioni, telco ed energia, oltre alla moda. Take operativo: la lista e costruita esattamente sui settori dove i budget esistono davvero.')

para('6. Concorrenza e acquisizione B2B', bold=True)
para('Esistono molte agenzie di produzione video tradizionali; il differenziale di AD Lab e l\'AI (qualita, tempi, costi) unita a un posizionamento alto. Sull\'acquisizione enterprise funziona l\'approccio ABM: pochi account selezionati, messaggio su misura, piu touch. Take operativo: meglio pochi appuntamenti ma molto qualificati che tanti generici.')
rule()

# ===== PARTE 2A =====
h1('PARTE 2A - MACRO A: Requisiti del target e ICP')
para('Queste risposte chiudono i criteri di selezione. Nella call di onboarding erano rimasti aperti due numeri (soglie fatturato/dipendenti) e un terzo parametro che volevi recuperare: qui li fissiamo.')

q('A1','Soglie minime di fatturato e dipendenti',
  'Quali sono le soglie minime di fatturato e di numero dipendenti sotto cui un\'azienda NON e in target? E qual era il terzo parametro che volevi indicarci (oltre a fatturato e dipendenti)?')
q('A2','Enterprise contro mid-large: dove concentrarci',
  'Nel Master abbiamo 132 enterprise/multinazionali (decisore difficile da raggiungere, ciclo lungo, tipo Haribo/Maserati) e 153 mid-large piu lavorabili (decisore raggiungibile, appuntamento piu rapido). Preferisci che partiamo dai mid-large per portare appuntamenti in fretta tenendo i giganti come lavoro di fondo, oppure vuoi spingere subito anche sui grandi brand?')
q('A3','Settori da tenere o tagliare',
  'La moda resta prioritaria. Gli altri cinque settori (food & beverage, automotive, retail/GDO, finance/assicurazioni, telco/energia) li teniamo tutti o ne escludiamo qualcuno perche meno coerente con cio che vendete?')
q('A4','Sotto-segmenti moda da privilegiare',
  'Dentro la moda, quali segmenti mettiamo come priorita A (verdi): lusso/alta moda, sportswear/outdoor, calzature/occhialeria/gioielli, beauty/cosmesi? In quale ordine?')
q('A5','Decision maker e multinazionali',
  'Confermi che gli interlocutori target sono il responsabile marketing + i C-level (CEO/CFO/CMO)? E che per le multinazionali puntiamo all\'entita italiana del gruppo (es. brand Italia)?')
q('A6','Dato investimento adv',
  'Avevi chiesto di sapere quanto le aziende investono in advertising. Lo abbiamo gia impostato come colonna "Evidenza investimento adv/contenuti" (campagne attive, spot, social, segnali di nuovo investimento). Va bene cosi o vuoi un formato/dettaglio diverso?')
rule()

# ===== PARTE 2B =====
h1('PARTE 2B - MACRO B: Lista, scouting e messaggio')
para('Abbiamo gia costruito e verificato 427 prospect (vedi Master_Eureweb_ADLab.xlsx). Per finalizzare lo scouting e lo script ci servono 5 risposte operative + 5 risposte "non-ovvie", pensate per scavare quello che altrimenti scopriremmo perdendo tempo.')

para('Blocco operativo - 5 domande', bold=True)
q('B1','Approvazione lista',
  'Scorrendo la lista (filtrabile per settore, priorita, "Quick win"), ci sono aziende che vuoi togliere subito per conflitto di interesse, perche non in linea, o perche non le ritieni adatte?')
q('B2','Aziende da non toccare',
  'Oltre ai vostri clienti gia attivi (gia esclusi dalla lista), ci sono prospect che state gia lavorando internamente e che NON dobbiamo contattare per non sovrapporci o bruciarli?')
q('B3','Casi citabili e showreel',
  'Quali clienti o lavori possiamo nominare in chiamata come riferimento, senza problemi di consenso (es. settore hotellerie, automotive, luxury)? E lo showreel di 20 secondi pubblicabile a che punto e: quando possiamo usarlo nell\'outreach?')
q('B4','Script: nostra versione poi confronto',
  'Ti prepariamo la nostra prima versione di script (senza prezzi, focalizzata sul fissare l\'appuntamento conoscitivo). Quando l\'hai validata, ci condividi i tuoi script interni per fare A/B test e migliorarli. Confermi questo flusso?')
q('B5','Canali e LinkedIn',
  'Lavoriamo su chiamata (centralino verso marketing) + LinkedIn + email. Confermi la combinazione? Avete profili LinkedIn (aziendali o personali) che possiamo usare per le richieste di collegamento ai decisori?')

para('Blocco non-ovvie - 5 domande dalla ricerca', bold=True)
q('B6','Differenziale difficile da replicare',
  'Cosa sa fare AD Lab che un concorrente fatica a replicare in 12 mesi (art director con esperienza nel lusso, ambassador di una piattaforma AI, resa fotorealistica sugli ambienti reali forniti dal cliente, altro)? Questo diventa il gancio della prima mail/chiamata a un brand nuovo.')
q('B7','Brand sognato e chi vi ha detto no',
  'Qual e il brand dei sogni che vorreste portare a casa? E c\'e un prospect che vi ha gia detto no, e per quale motivo (cosi non ripetiamo l\'approccio sbagliato)?')
q('B8','Obiezioni piu frequenti',
  'Quando proponete shooting AI, quali sono le 2-3 obiezioni che sentite piu spesso (sembra finto? perdiamo controllo del brand? il nostro target non lo accetterebbe?)? Cosi le anticipiamo nello script.')
q('B9','Leva di valore in chiamata (senza prezzo)',
  'Possiamo usare come argomento "stessa qualita di uno shooting tradizionale, con tempi e costi piu efficienti" e la rapidita di produzione, oppure preferisci che spingiamo su un\'altra angolazione (es. versatilita su tutti i touchpoint, qualita creativa, velocita di adattamento)?')
q('B10','Barter / cambio merce',
  'Confermiamo che il cambio merce (DigitalKPI) NON si nomina in chiamata e resta una tua leva in fase di chiusura, giusto? O in alcuni casi vuoi che lo accenniamo?')
rule()

# ===== PARTE 3 =====
h1('PARTE 3 - KPI e modalita di lavoro')
para('Le 4 risposte qui sotto chiudono la parte operativa: misuriamo il successo, fissiamo la cadenza e calibriamo sul calendario estivo.')

q('K1','Numero di appuntamenti al mese',
  'Vogliamo allinearci sul numero giusto: a contratto risultano almeno 10 appuntamenti in target al mese, e ci avevi detto che ne vorresti anche di piu. Confermi 10/mese come minimo, con la regola "meglio pochi ma in target e di qualita che tanti generici"?')
q('K2','Cosa vuoi vedere a fine luglio',
  'Quali numeri ti fanno dire "ha funzionato"? Es: X appuntamenti in target fissati, Y aziende realmente interessate, pipeline generata. Qual e per te la soglia di successo realistica per giugno-luglio?')
q('K3','Cadenza di review',
  'Preferisci un punto operativo settimanale (es. 30 minuti il venerdi) per aggiustare il tiro, oppure un report scritto periodico + una call mensile?')
q('K4','Calendario estivo',
  'Confermi campagna su giugno e luglio con stop ad agosto? Quando ripartiamo a settembre, cosi calibriamo le ultime settimane di luglio e teniamo pronta la lista per la ripartenza?')
rule()

para('Le risposte ci permettono di chiudere i requisiti, mandarti la lista definitiva da approvare e partire subito con lo scraping dei contatti diretti sui prospect "Quick win" (i mid-large in target piu raggiungibili). Grazie Giulia.', italic=True, size=9.5)

out='/Users/simocors/Desktop/telesales/eureweb/Questionario_Eureweb_ADLab.docx'
doc.save(out)
print('SALVATO', out)
