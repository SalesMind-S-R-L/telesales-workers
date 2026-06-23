#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Dossier di mercato Eureweb / AD Lab - contenuti foto/video AI per grandi brand.
Dati verificati multi-fonte (giugno 2026). Niente emoji. Stesso impianto del dossier IMEON.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY = RGBColor(0x1F,0x38,0x64)
doc = Document()
st = doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(10.5)

def h1(t):
    p=doc.add_heading(t,level=1)
    for r in p.runs: r.font.color.rgb=NAVY
def h2(t):
    p=doc.add_heading(t,level=2)
    for r in p.runs: r.font.color.rgb=NAVY
def para(t,bold=False,italic=False,size=None):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=bold; r.italic=italic
    if size: r.font.size=Pt(size)
    return p
def bullet(t):
    doc.add_paragraph(t,style='List Bullet')
def table(headers, rows):
    t=doc.add_table(rows=1, cols=len(headers)); t.style='Light Grid Accent 1'
    hc=t.rows[0].cells
    for i,hh in enumerate(headers):
        hc[i].text=''; run=hc[i].paragraphs[0].add_run(hh); run.bold=True; run.font.size=Pt(9)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=''; rr=cells[i].paragraphs[0].add_run(str(v)); rr.font.size=Pt(9)
    return t

# ---------------- COVER ----------------
title=doc.add_paragraph(); title.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=title.add_run('DOSSIER DI MERCATO'); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=NAVY
sub=doc.add_paragraph(); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=sub.add_run('Contenuti foto/video AI high-ticket per grandi brand italiani'); r.font.size=Pt(14); r.font.color.rgb=NAVY
s2=doc.add_paragraph(); s2.alignment=WD_ALIGN_PARAGRAPH.CENTER
s2.add_run('Per Eureweb / unit "AD Lab" - shooting foto/video generati con intelligenza artificiale').italic=True
s3=doc.add_paragraph(); s3.alignment=WD_ALIGN_PARAGRAPH.CENTER
s3.add_run('Aggiornamento: giugno 2026  |  Settori: Automotive, Food & Beverage, Retail/GDO, Fashion/Lifestyle, Finance/Insurance, Telecomunicazioni').font.size=Pt(9)

para('Nota metodologica - i due "numeri" del mercato pubblicitario italiano. Coesistono due misure diverse e spesso confuse: Nielsen Ad Intel rileva un perimetro a panel (chiude il 2025 a ~9,6 mld EUR); l\'Osservatorio Internet Media del Politecnico di Milano misura il mercato totale, incluso tutto il digitale (2025 a ~11,8 mld EUR). Entrambi sono corretti: misurano cose diverse. Nel dossier si riporta sempre la fonte accanto al dato.',italic=True,size=9)

# ---------------- 1 ----------------
h1('1. Mercato pubblicitario italiano 2024-2025')
table(['Indicatore','Dato','Fonte'],[
 ['Mercato totale 2024','11,1 mld EUR (+8% vs 2023)','Osservatorio Internet Media PoliMi'],
 ['Mercato totale 2025','11,8 mld EUR (+5%)','Osservatorio Internet Media PoliMi'],
 ['Raccolta Nielsen Ad Intel 2024','~9,5 mld EUR (+3,8%)','Nielsen (Engage/Primaonline)'],
 ['Raccolta Nielsen Ad Intel 2025','9,6 mld EUR (+0,4%)','Nielsen (Engage)'],
])
h2('Split per mezzo (mercato totale, fonte PoliMi)')
table(['Mezzo','Quota / valore','Andamento'],[
 ['Internet / Digital','~50% (5,4 mld 2024; 6,2 mld e 53% nel 2025)','+10/12%'],
 ['TV','~35% (circa 4,3 mld 2024)','+6%'],
 ['Out of Home','6-7%','+13% (2024)'],
 ['Stampa (quotidiani+periodici)','5-6%','-2/-8%'],
 ['Radio','4%','+8% (2024)'],
])
h2('Video advertising - il dato piu rilevante per AD Lab')
bullet('2024: incidenza del video sul totale advertising ~43%, mercato ~4,7 mld EUR. (PoliMi / Confindustria Radio TV)')
bullet('2025: video online 2,9 mld EUR (+16%), pari al 41% dell\'internet advertising; la TV connessa cresce a ~832 mln EUR (+19%). (Osservatorio PoliMi)')
para('Lettura per la vendita: il video e la voce che cresce di piu ed e proprio il terreno dove l\'AI generativa abbatte i costi. E l\'angolo d\'attacco piu forte.',bold=True)
para('Divergenza Nielsen 2025: Nielsen vede la TV in calo (-1,8%) e il mercato piatto (+0,4%), mentre PoliMi vede +5% trainato dal digitale (perimetri diversi). Usare PoliMi per "il mercato cresce", Nielsen per "i mezzi tradizionali soffrono".',italic=True,size=9)

# ---------------- 2 ----------------
h1('2. Spesa per settore e top advertiser')
h2('Settori che spendono di piu (H1 2024, Nielsen Ad Intel)')
table(['Settore / Categoria','Spesa','Note'],[
 ['Automotive','~239 mln EUR (top categoria)','Pressione su nuovi modelli e transizione elettrica'],
 ['Top 5 categorie aggregate','~2,63 mld EUR = 44,9% del totale','Alimentari, Distribuzione/GDO, Automotive, Farmaceutici, Abitazione'],
 ['E-commerce e servizi online','~156 mln EUR','-'],
 ['Farmaceutici OTC','~135 mln EUR','-'],
 ['Bevande analcoliche','~92 mln EUR','-'],
])
para('Top advertiser singoli in Italia (H1 2024, Nielsen): 1) Procter & Gamble, 2) Ferrero, 3) Volkswagen, 4) L\'Oreal, 5) Unilever, 6) Fiat. (FoodAffairs/Nielsen, Mark Up)')
h2('Andamento per settore - i 6 target di AD Lab')
table(['Settore target','Andamento','Lettura commerciale'],[
 ['Automotive','2024 +14,3%; 2025 -5%','Budget alti ma sotto pressione costi: AI come leva di efficienza'],
 ['Food & Beverage','Top categoria; alimentari +3,7%, bevande/alcoolici +6,7% (2025)','Settore piu ricco e in crescita, alta frequenza di contenuti: ideale'],
 ['Retail / GDO','2024 +7,8%; 2025 -10,2%','Sotto pressione marginalita: messaggio "stesso impatto, meno budget"'],
 ['Finance / Insurance','2025 +11,7% (miglior crescita)','In espansione adv, forte su brand/trust: opportunita calda'],
 ['Telecomunicazioni','2024 -9,7%; 2025 debole','Tagliano budget: leva AI = efficienza pura'],
 ['Fashion / Lifestyle','Alta stagionalita, alto volume shooting','Settore AI-friendly per eccellenza: cataloghi, lookbook, varianti'],
])
para('Insight: i 6 target si dividono in due gruppi. Crescenti (Food & Beverage, Finance/Insurance): vendere "fai di piu col budget che gia cresce". Sotto pressione (Automotive 2025, Retail/GDO, Telco): vendere "stesso risultato, costo abbattuto". Fashion/Lifestyle e il caso d\'uso AI per eccellenza.',bold=True)

# ---------------- 3 ----------------
h1('3. AI nella produzione di contenuti e advertising')
h2('Dimensione e crescita del mercato AI in Italia')
table(['Indicatore','Dato','Fonte'],[
 ['Mercato AI Italia 2025','1,8 mld EUR (+50% vs 2024)','Osservatorio AI PoliMi'],
 ['Mercato AI Italia 2024','~1,2 mld EUR (+58% su 2023)','Osservatorio AI PoliMi'],
 ['Quota GenAI / progetti ibridi','46% del mercato','Osservatorio AI PoliMi'],
])
h2('Adozione GenAI nelle imprese (chiave per il pitch)')
table(['Indicatore','Dato','Fonte'],[
 ['Grandi imprese con >=1 licenza GenAI (2025)','84% (+31 punti vs 2024)','Osservatorio AI PoliMi'],
 ['Grandi imprese con >=1 progetto AI (2025)','71% (era 59% nel 2024)','Osservatori PoliMi (LineaEDP)'],
 ['Usa GenAI per creazione contenuti','~50% degli intervistati','Osservatori (Channeltech)'],
 ['Imprese 10+ dip. con AI (totale, ISTAT/Eurostat)','8,2% (sotto media UE 13,5%)','LineaEDP/Eurostat'],
])
para('Convivono due fotografie. Tra le grandi imprese l\'adozione e altissima (84% ha licenze GenAI). Sul totale aziende (incluse PMI) la penetrazione e ancora bassa (8,2%). Per AD Lab, che vende a enterprise, conta il primo numero: il decisore enterprise ha gia l\'AI in casa. Non bisogna convincere che "l\'AI esiste", ma posizionarsi come specialista high-end.',bold=True)
h2('GenAI nel video advertising (dato piu forte del dossier)')
table(['Indicatore','Dato','Fonte'],[
 ['Video advertising globali fatti/migliorati con GenAI','30% nel 2025 (era 22% nel 2024)','IAB (AgendaDigitale/IBM)'],
 ['Caso Coca-Cola','Spot di Natale 2024 e 2025 con AI; per il 2025 70.000 videoclip generati da ~100 persone+specialisti','Mediatrends/AOL'],
 ['Caso Italia - RAI','Sanremo Giovani 2026: spot promozionale interamente in GenAI','Intervista.it'],
 ['Caso Italia - Kortocircuito','Spot "full AI casting" in onda su emittenti nazionali','Kortocircuito'],
])
h2('Costo tradizionale vs AI - benchmark')
table(['Voce','Costo','Fonte'],[
 ['Shooting fotografico pro','100-600 EUR a sessione; fotografo 400-1.600 EUR/giorno','Cromie/Concept'],
 ['Spot pubblicitario (produzione)','da ~3.000 a 10.000+ EUR per video professionale','Kortocircuito/ProntoPro'],
 ['Campagna TV completa (produzione + spazi)','250.000-300.000 EUR','Kortocircuito'],
 ['Riferimento citato dal cliente (video tradizionale)','~50.000 EUR (ordine di grandezza coerente con spot alta fascia)','Benchmark di settore'],
 ['Risparmio con AI','fino a -60% costo e +50% velocita','Stime di settore (vendor/agenzie)'],
])
para('Da maneggiare con onesta: il "-60% / +50%" proviene da fonti di vendor/agenzie AI, non da uno studio indipendente. Il riferimento "video tradizionale ~50.000 EUR" e plausibile (in linea con i benchmark di produzione spot di fascia alta) ma non riconducibile a una singola fonte ufficiale. In chiamata si puo citare l\'ordine di grandezza, non spacciarlo per dato Nielsen.',italic=True,size=9)

# ---------------- 4 ----------------
h1('4. Chi decide i contenuti e la creativita nei grandi brand')
table(['Ruolo','Cosa decide','Rilevanza per AD Lab'],[
 ['CMO / Direttore Marketing','Strategia, budget annuale, scelta partner; risponde a CEO/CdA','Decisore economico per contratti high-ticket. E a lui/lei che serve arrivare'],
 ['Brand Manager / Marketing Manager','Esecuzione campagne, gestione fornitori creativi','Influenzatore forte, spesso primo contatto'],
 ['Head of Content / Content Manager','Pianificazione e produzione contenuti, volumi, calendario','Sente il dolore di costi/tempi: ottimo champion interno'],
 ['Digital Marketing Manager','Performance, canali digitali, asset per adv','Champion sui formati video/social ad alto volume'],
 ['Creative / Media Procurement','Negoziazione, gare, condizioni economiche','Gatekeeper nelle enterprise strutturate; entra a valle'],
])
h2('Struttura del processo d\'acquisto (enterprise)')
bullet('Il Direttore Marketing guida la scelta dei partner; a lui riportano Digital Marketing Manager e PR/Comunicazione; interagisce con CEO/CdA.')
bullet('Nelle grandi aziende l\'acquisto di servizi creativi passa spesso da gare/RFP e dal coinvolgimento del Procurement, con ciclo decisionale lungo e multi-stakeholder (il marketing decide "cosa", il procurement "a quali condizioni").')
bullet('Tempi tipici: i cicli B2B enterprise sono lunghi e multi-touch (settimane/mesi), con piu decisori coinvolti e nurturing prolungato.')
para('Implicazione operativa: in chiamata di setting non si vende e non si negozia col procurement. Si fissa un incontro conoscitivo con il decisore marketing (CMO/Direttore/Head of Content). Il primo "si" da ottenere e l\'appuntamento, non l\'ordine.',bold=True)

# ---------------- 5 ----------------
h1('5. Concorrenza')
table(['Player','Profilo','Fonte'],[
 ['AQuest','Creative Production & Technology Company, dal 2019 nel gruppo WPP, Gold Partner Microsoft','Sortlist/web'],
 ['ORBIS Production','Casa di produzione + agenzia creativa integrata, brand globali, 16+ anni','Bliss/web'],
 ['Kortocircuito','Agenzia di produzione video con AI; primo spot "full AI casting" su TV nazionale','Kortocircuito'],
 ['Studi virtual production (Milano)','65% delle case di produzione ha investito in AI/virtual production negli ultimi 18 mesi','Mediatica'],
 ['Network globali (WPP, Publicis, Omnicom IT)','Grandi gruppi che integrano GenAI nella filiera','Settore'],
 ['Piattaforme AI (Runway, Sora, Adobe Firefly)','Concorrenza "fai-da-te" lato cliente','Settore'],
])
para('Posizionamento per AD Lab: il mercato si divide in (a) grandi network/case di produzione che aggiungono AI alla filiera, (b) piattaforme self-service economiche, (c) poche realta AI-first specializzate. Spazio bianco: "AI-first high-end per enterprise" - qualita da grande produzione con economics e velocita AI. La minaccia non e il concorrente diretto ma il "facciamolo internamente con Sora/Firefly". Il pitch deve spostare la conversazione da "strumento" a "risultato + responsabilita + qualita brand-safe".',bold=True)
para('Benchmark prezzo: produzione spot alta fascia/campagna ~50.000-300.000 EUR; AI fino a -60% costo. Il prezzo AD Lab (da 15.000 EUR) si colloca sopra il self-service e sotto la grande produzione tradizionale: posizionamento "premium ma efficiente" difendibile.')

# ---------------- 6 ----------------
h1('6. Leve di acquisizione B2B')
table(['Leva','Come funziona','Evidenza'],[
 ['ABM (Account-Based Marketing)','Pochi account enterprise, ricerca account-specific, asset personalizzati, multi-touch','Caso Snowflake: +38% account con meeting, 3x meeting rate sui tier alti'],
 ['Outbound mirato','SDR/setter su lista pre-qualificata di decisori marketing; obiettivo = appuntamento','UnboundB2B/Ironpaper'],
 ['Case study & social proof','Mostrare risultati reali (es. spot AI in onda): driver di fiducia n.1','Pedowitz/Understory'],
 ['Thought leadership / contenuti','Demand gen organica + ABM per convertire account ad alto fit','Gripped'],
 ['Referral / portfolio','Brand-name nel portfolio abbassa il rischio percepito dal CMO','Settore B2B'],
 ['Demo "show, don\'t tell"','Per un servizio AI visivo una demo concreta vale piu di mille slide','Settore'],
])
para('Sintesi: per AD Lab la combinazione vincente e outbound/ABM mirato sui 6 settori -> appuntamento conoscitivo -> demo + case study in incontro. La chiamata serve solo ad aprire la porta.',bold=True)

# ---------------- 7 ----------------
h1('7. Takeaway operativi (insight -> uso in chiamata di setting)')
para('Regola ferma: in chiamata NON si fanno prezzi. Obiettivo unico = fissare l\'appuntamento conoscitivo con il decisore marketing.',bold=True)
tk=[
 ('Video adv ~2,9 mld EUR 2025 (+16%, 41% del digitale) - PoliMi','"Il video e il formato che cresce di piu nel mercato pubblicitario. Come state gestendo la produzione dei vostri contenuti video? Vale 15 minuti."'),
 ('84% delle grandi imprese italiane ha gia licenze GenAI - Oss. AI PoliMi','"So che realta come la vostra usano gia l\'AI internamente. Noi la applichiamo alla produzione creativa per i brand: confrontiamoci in una call."'),
 ('30% degli spot video nel mondo gia fatto/migliorato con GenAI (era 22%) - IAB','"Quasi un terzo del video advertising globale passa per l\'AI. Vorrei mostrarvi cosa significa per un brand del vostro livello: 20 minuti?"'),
 ('Coca-Cola: 70.000 videoclip AI per un solo spot - Mediatrends','"I grandi brand producono migliaia di varianti con l\'AI. Il punto non e \'se\' ma \'come farlo bene\'. Vi va di vedere come lo facciamo noi?"'),
 ('Automotive top spender (~239 mln) ma 2025 -5% - Nielsen','(auto) "Il vostro settore investe piu di tutti, ma i budget sono sotto pressione. Vi mostro come ottenere lo stesso impatto ottimizzando la produzione."'),
 ('Finance/Insurance miglior crescita adv 2025 (+11,7%) - Nielsen','(finance) "Il vostro settore e tra i pochi che aumentano gli investimenti: e il momento giusto per fare di piu col budget. Fissiamo un incontro?"'),
 ('Retail/GDO e Telco tagliano i budget (-10% / -10%) - Nielsen','(retail/telco) "C\'e massima attenzione ai costi: proprio per questo voglio presentarvi un modo per produrre contenuti di alto livello in meno tempo."'),
 ('Produzione spot alta fascia: decine di migliaia di EUR; AI riduce molto tempi/costi - settore','"Senza numeri al telefono perche dipendono dal progetto: la differenza in tempi e flessibilita e enorme. Preferisco mostrarvelo dal vivo."'),
 ('Mercato AI Italia +50%/anno; a Milano 65% delle case di produzione ha investito in AI - PoliMi/Mediatica','"Il mercato si muove molto in fretta su questo. Vi do una lettura di dove sta andando la produzione contenuti, anche se non lavoriamo insieme."'),
 ('Decide il marketing; il procurement entra dopo - analisi B2B','"Per non farvi perdere tempo: la persona giusta per una panoramica e chi segue marketing e contenuti. Parlo con la persona corretta?"'),
]
table(['Insight chiave (dato verificato)','Come usarlo in chiamata'],tk)

# ---------------- FONTI ----------------
h1('Fonti')
para('Mercato pubblicitario',bold=True)
for u in ['Nielsen 2025: engage.it/dati-e-ricerche/nielsen-investimenti-pubblicitari-2025.aspx',
 'Nielsen 2024: engage.it/dati-e-ricerche/nielsen-dati-pubblicita-2024.aspx',
 'Nielsen 2024 (Primaonline): primaonline.it/2025/02/12/433104/',
 'Osservatorio Internet Media PoliMi: osservatori.net/comunicato/internet-media/internet-advertising-italia-in-crescita/',
 'PoliMi 11,6-11,8 mld 2025 (ItaliaOggi): italiaoggi.it (mercato pubblicitario PoliMi 2025)',
 'IAB Italia 2023: iab.it/il-mercato-pubblicitario-italiano-nel-2023-vale-10-2-miliardi-di-euro/',
 'Raccolta TV 4,3 mld 2024 (ADC Group): adcgroup.it (raccolta pubblicitaria TV 2024)',
 'Confindustria Radio TV (video 43%): confindustriaradiotv.it']:
    bullet(u)
para('Spesa per settore e top advertiser',bold=True)
for u in ['Top 10 inserzionisti + categorie (FoodAffairs/Nielsen): foodaffairs.it/2024/10/24/classifica-dei-primi-10-inserzionisti-pubblicitari-in-italia',
 'Top spender Nielsen (Brand News): brand-news.it/intelligence/dati/pubblicita-ecco-i-top-spender-in-italia-secondo-nielsen-ad-intel/',
 'Settori che investono di piu (Influent People): influentpeople.it',
 'Ferrero top spender (Mark Up): mark-up.it/le-10-aziende-che-investono-di-piu-in-pubblicita-ferrero-al-top/']:
    bullet(u)
para('AI, produzione contenuti e mercato AI',bold=True)
for u in ['Osservatorio AI PoliMi 2025 (1,8 mld, 84% licenze): osservatori.net/comunicato/artificial-intelligence/intelligenza-artificiale-italia/',
 'AI4Business: ai4business.it (mercato AI Italia +50% 2025)',
 'AI 2024 1,2 mld +58% (Aimage): aimage.it/2025/02/07/',
 'Adozione GenAI grandi imprese (LineaEDP): lineaedp.it/report/intelligenza-artificiale-e-genai-adozione-in-crescita-nel-2024/',
 'GenAI motore di crescita (Channeltech): channeltech.it/2025/10/24/',
 'Adozione GenAI globale x5 (Data Manager): datamanager.it/2025/10/']:
    bullet(u)
para('AI nel video advertising / casi brand',bold=True)
for u in ['IAB 30% video con GenAI; casi Google/Coca-Cola (Agenda Digitale): agendadigitale.eu/cultura-digitale/spot-fatti-con-lai-cosa-funziona-e-cosa-no-i-casi-google-e-coca-cola/',
 'Coca-Cola spot AI 2025 (Mediatrends): mediatrends.it/coca-cola-ia-spot-natale-mediatrends/',
 'Coca-Cola GenAI creativa (Intervista.it): intervista.it/aziende/coca-cola-nel-2025-punta-sempre-di-piu-sulla-gen-ai-creativa/',
 'AI-generated advertising 2025 (IBM): ibm.com/think/news/ai-generated-advertising-2025',
 'Coca-Cola 70.000 clip (AOL): aol.com/articles/coca-cola-doubles-down-ai-220439963.html']:
    bullet(u)
para('Costi produzione tradizionale vs AI',bold=True)
for u in ['Costo shooting (Cromie): comunicazione.agency/shooting-fotografico-come-funziona-quanto-costa/',
 'Costo shooting (Concept): conceptsnc.com/it/costo-shooting-fotografico/',
 'AI vs photoshoot (Prime AI): prime-ai.com/en/media/cost-comparison-budget-photoshoots-vs-ai-photoshoot-c/',
 'Costo spot TV (ManagerAds): managerads.it/costi-produzione-spot-tv/',
 'Costo spot (Kortocircuito): kortocircuito.com/magazine/quanto-costa-uno-spot-pubblicitario-in-tv/']:
    bullet(u)
para('Concorrenza, decisori e acquisizione B2B',bold=True)
for u in ['Agenzie video Italia (Bliss): blissagency.it/uncategorized/le-10-migliori-agenzie-di-produzione-video-in-italia/',
 'Agenzie creative (Sortlist): sortlist.it/s/creativo/italia-it',
 'IA generativa produzione audiovisiva + 65% Milano (Mediatica): mediaticanetwork.com/magazine/lia-generativa-ridefinisce-la-produzione-audiovisiva',
 'Chi decide il marketing (Vittoria Comunica): vittoriacomunica.it/blog/chi-decide-marketing-azienda/',
 'Ruolo CMO / Direttore Marketing: marketingstrategy.solutions',
 'ABM enterprise + caso Snowflake (FieldTrip): fieldtrip.agency/post/best-abm-agencies-for-enterprise-b2b',
 'B2B creative agencies (Pedowitz): pedowitzgroup.com/blog/top-9-b2b-creative-agencies-for-enterprise-stories']:
    bullet(u)

# ---------------- CAVEAT ----------------
h1('Caveat finali sull\'affidabilita dei dati')
bullet('Verificati su fonti primarie/multiple: mercato pubblicitario (Nielsen + PoliMi), spesa per settore, mercato AI Italia, adozione GenAI grandi imprese, % video adv con GenAI (IAB), casi Coca-Cola.')
bullet('Da usare con cautela (fonti vendor/agenzie): il "-60% costo / +50% tempo" dell\'AI e il riferimento "video tradizionale ~50.000 EUR" sono ordini di grandezza di settore plausibili, non dati ufficiali. Presentarli come "stime/benchmark di settore", non come dati certificati.')
bullet('Doppio standard di mercato: ricordare sempre la distinzione Nielsen (panel, ~9,6 mld) vs PoliMi (mercato totale, ~11,8 mld) per evitare incongruenze.')

doc.add_paragraph()
para('Documento di lavoro interno Telesales per il cliente Eureweb / AD Lab. I dati hanno data e fonte indicate; verificare gli aggiornamenti prima di usarli in contesti ufficiali.',italic=True,size=8)

out='/Users/simocors/Desktop/telesales/eureweb/Dossier_Mercato_Eureweb.docx'
doc.save(out)
print('SALVATO',out)
